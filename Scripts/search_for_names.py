import os
from docx import Document

# Função para ler o conteúdo de um arquivo .docx, incluindo tabelas
def read_docx_with_tables(file_path):
    doc = Document(file_path)
    full_text = []
    
    # Adiciona o texto fora das tabelas
    for para in doc.paragraphs:
        full_text.append(para.text)
    
    # Adiciona o texto dentro das tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)
    
    return '\n'.join(full_text)

# Função para procurar nomes dentro de arquivos Word em uma pasta
def find_names_in_word_files(directory, names):
    results = {}
    
    # Percorre os arquivos na pasta
    for filename in os.listdir(directory):
        if filename.endswith(".docx"):
            file_path = os.path.join(directory, filename)
            try:
                content = read_docx_with_tables(file_path)
                for name in names:
                    if name.lower() in content.lower():
                        if filename not in results:
                            results[filename] = []
                        results[filename].append(name)
            except Exception as e:
                print(f"Erro ao ler o arquivo {filename}: {e}")
    
    return results

# Diretório onde estão os arquivos Word
directory = r'C:\Users\joao.ribeiro\Documents\GPT'

# Lista de nomes a serem procurados (Substitua com os nomes que deseja procurar)
names_to_find = [
    "Henrique",
    "Rodney",
    "Sebastião",
]

# Chama a função para procurar os nomes
found_names = find_names_in_word_files(directory, names_to_find)

# Exibe os resultados
if found_names:
    for filename, found in found_names.items():
        print(f"Nomes encontrados no arquivo {filename}: {', '.join(found)}")
else:
    print("Nenhum nome foi encontrado nos arquivos.")