import os
import docx
import pandas as pd

def extract_word_data(file_path):
    doc = docx.Document(file_path)
    
    # Variáveis para armazenar os dados
    numero_fsa = None
    data_acesso_inicio = None
    data_acesso_termino = None
    horario_acesso = None
    pessoas = []

    # Laço para processar os parágrafos
    for para in doc.paragraphs:
        text = para.text.strip()
        # Aceitar diferentes variações para o número da FSA
        if text.startswith("FSA N") or text.startswith("FSA Nº") or text.startswith("FSA N°"):
            numero_fsa = text
        elif "Início:" in text and "Término:" in text and "Horário:" in text:
            # Extraindo data de início, término e horário
            data_acesso_inicio = text.split("Início:")[-1].split("Término:")[0].strip()
            data_acesso_termino = text.split("Término:")[-1].split("Horário:")[0].strip()
            horario_acesso = text.split("Horário:")[-1].strip()

    # Laço para processar as tabelas
    for table in doc.tables:
        for i, row in enumerate(table.rows[1:]):  # Ignorar o cabeçalho
            cells = [cell.text.strip() for cell in row.cells]
            pessoa = {}
            
            if len(cells) > 0:
                pessoa["Nome Completo"] = cells[0]
            if len(cells) > 1:
                pessoa["Filiação"] = cells[1]
            if len(cells) > 2:
                pessoa["País, Cidade e Data de Nascimento"] = cells[2]
            if len(cells) > 3:
                pessoa["Cidadania"] = cells[3]
            if len(cells) > 4:
                pessoa["Outras Nacionalidades"] = cells[4]
            if len(cells) > 5:
                pessoa["Endereço Residencial"] = cells[5]
            if len(cells) > 6:
                pessoa["Governo/Empresa/Entidade"] = cells[6]
            if len(cells) > 7:
                pessoa["Profissão Atual"] = cells[7]
            if len(cells) > 8:
                pessoa["Cargo/Função Atual"] = cells[8]
            if len(cells) > 9:
                pessoa["Endereço Funcional"] = cells[9]

            pessoas.append(pessoa)

    # Retornar os dados extraídos
    return {
        "Número da FSA": numero_fsa,
        "Início Acesso": data_acesso_inicio,
        "Término Acesso": data_acesso_termino,
        "Horário Acesso": horario_acesso,
        "Pessoas": pessoas
    }

def sanitize_sheet_name(sheet_name):
    """Remove caracteres inválidos de nomes de planilhas."""
    invalid_chars = ['\\', '/', '*', '[', ']', ':', '?', '"']
    for char in invalid_chars:
        sheet_name = sheet_name.replace(char, '_')  # Substitui por _
    return sheet_name

def process_multiple_files(folder_path, output_excel_path):
    # Criar um ExcelWriter para salvar várias planilhas
    with pd.ExcelWriter(output_excel_path, engine='openpyxl') as writer:
        # Laço por todos os arquivos .docx na pasta
        for filename in os.listdir(folder_path):
            if filename.endswith(".docx") and not filename.startswith("~$"):  # Ignora arquivos temporários
                file_path = os.path.join(folder_path, filename)
                # Extrair dados de cada arquivo Word
                dados = extract_word_data(file_path)

                # Verifica se o número da FSA foi encontrado
                if dados["Número da FSA"] is None:
                    print(f"Número da FSA não encontrado em {filename}. Usando nome padrão.")
                    sheet_name = sanitize_sheet_name(f"Planilha_{filename}")
                else:
                    # Nome da planilha será o número da FSA, sanitizado
                    sheet_name = sanitize_sheet_name(dados["Número da FSA"].replace("FSA Nº", "").replace("FSA N°", "").strip())
                
                # Criar um DataFrame para os dados das pessoas
                df = pd.DataFrame(dados["Pessoas"])
                
                # Adicionar os metadados (Número da FSA, datas) no início da planilha
                df_meta = pd.DataFrame({
                    "Número da FSA": [dados["Número da FSA"]],
                    "Início Acesso": [dados["Início Acesso"]],
                    "Término Acesso": [dados["Término Acesso"]],
                    "Horário Acesso": [dados["Horário Acesso"]]
                })
                
                # Concatenar os metadados com os dados das pessoas
                final_df = pd.concat([df_meta, df], axis=1)
                
                # Escrever os dados na planilha
                final_df.to_excel(writer, sheet_name=sheet_name, index=False)

# Exemplo de uso
folder_path = r"FSAs"  # Atualize para o caminho correto
output_excel_path = r"output.xlsx"  # Atualize para o caminho desejado
process_multiple_files(folder_path, output_excel_path)
